import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "prompt-compression-whitepaper.md"
OUTPUT = ROOT / "Measuring_the_Prompt_Compression_Fidelity_Tradeoff.docx"
CHART = ROOT / "compression_reduction_by_workload.png"
CLIENT_SCALE_IMAGE = Path(r"C:\Users\troym\AppData\Local\Temp\codex-clipboard-ed5281fd-d359-4993-945c-12818589e96b.png")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203748"
GOLD = "A06A18"
MUTED = "666666"
LIGHT = "F4F6F9"
GRID = "C9D2DC"


def set_font(run, name="Calibri", size=11, color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_font(run, size=9, color=MUTED)


def paragraph_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def new_numbering_instance(doc, start=1):
    style = doc.styles["List Number"]._element
    base_node = style.find("./" + qn("w:pPr") + "/" + qn("w:numPr") + "/" + qn("w:numId"))
    base_num_id = base_node.get(qn("w:val"))
    numbering = doc.part.numbering_part.element
    base_num = next(node for node in numbering.findall(qn("w:num")) if node.get(qn("w:numId")) == base_num_id)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    node = num_pr.find(qn("w:numId"))
    if node is None:
        node = OxmlElement("w:numId")
        num_pr.append(node)
    node.set(qn("w:val"), str(num_id))


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def add_rich_text(paragraph, text):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            set_font(paragraph.add_run(text[pos:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            set_font(paragraph.add_run(token[1:-1]), italic=True)
        pos = match.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]))


def add_table(doc, lines):
    rows = [[c.strip() for c in line.strip().strip("|").split("|")] for line in lines]
    rows = [rows[0]] + rows[2:]
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    if cols == 6:
        widths = [2880, 720, 1440, 1440, 1260, 1620]
    elif cols == 4:
        widths = [3300, 2020, 2020, 2020]
    elif cols == 3:
        widths = [4200, 2580, 2580]
    else:
        widths = [9360 // cols] * cols
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        set_cant_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            add_rich_text(p, value)
            for run in p.runs:
                set_font(run, size=9.2, bold=(r_idx == 0), color=(INK if r_idx == 0 else "000000"))
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def create_chart():
    labels = ["FocusFit\nbefore", "FocusFit\nafter", "DeliveryTower\nbefore", "DeliveryTower\nafter"]
    values = [5.369, 5.278, 13.718, 13.675]
    colors = ["#7397B8", "#2E74B5", "#C49A58", "#A06A18"]
    width, height = 1500, 650
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype("arial.ttf", 30)
    small = ImageFont.truetype("arial.ttf", 25)
    left, right, top, bottom = 145, 50, 55, 125
    plot_w, plot_h = width-left-right, height-top-bottom
    for tick in (0, 4, 8, 12, 16):
        y = top + plot_h - int(plot_h * tick / 16)
        draw.line((left, y, width-right, y), fill="#DDE3E9", width=2)
        draw.text((80, y-15), str(tick), fill="#555555", font=small)
    draw.line((left, top, left, top+plot_h), fill="#777777", width=2)
    draw.line((left, top+plot_h, width-right, top+plot_h), fill="#777777", width=2)
    slot = plot_w / 4
    bar_w = 155
    for i, (label, value, color) in enumerate(zip(labels, values, colors)):
        cx = left + slot * (i + 0.5)
        bar_h = plot_h * value / 16
        x0, y0 = int(cx-bar_w/2), int(top+plot_h-bar_h)
        draw.rounded_rectangle((x0, y0, int(cx+bar_w/2), top+plot_h), radius=8, fill=color)
        value_text = f"{value:.2f}%"
        bbox = draw.textbbox((0, 0), value_text, font=font)
        draw.text((cx-(bbox[2]-bbox[0])/2, y0-42), value_text, fill="#203748", font=font)
        for j, part in enumerate(label.split("\n")):
            bbox = draw.textbbox((0, 0), part, font=small)
            draw.text((cx-(bbox[2]-bbox[0])/2, top+plot_h+14+j*29), part, fill="#444444", font=small)
    image.save(CHART)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(3)
    set_font(header.add_run("FOCUSEDOBJECTIVE  /  RESEARCH PREPRINT"), size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_after = Pt(0)
    add_page_field(footer)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("RESEARCH WHITEPAPER"), size=10, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("Measuring the Prompt\nCompression–Fidelity Tradeoff"), size=30, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(38)
    set_font(p.add_run("Initial evidence from production-like agent traces\nand a research design for detecting degradation"), size=15, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("FocusedObjective / PromptCompression"), size=12, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Preprint v0.1  |  July 12, 2026"), size=10, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(84)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Open source: github.com/FocusedObjective/PromptCompression\nHosted service: compress.usagetap.com"), size=9, color=MUTED)
    doc.add_page_break()


def build():
    create_chart()
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    doc.core_properties.title = "Measuring the Prompt Compression-Fidelity Tradeoff"
    doc.core_properties.subject = "Initial benchmark findings and semantic evaluation design"
    doc.core_properties.author = "FocusedObjective"
    add_cover(doc)

    in_quote = False
    table_lines = []
    inserted_chart = False
    skip_front = True
    number_list_id = None
    for line in lines:
        if skip_front:
            if line.strip() == "## Abstract":
                skip_front = False
            else:
                continue
        if table_lines and not line.startswith("|"):
            add_table(doc, table_lines)
            table_lines = []
        if line.startswith("|"):
            table_lines.append(line)
            continue
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            number_list_id = None
            p = doc.add_paragraph(stripped[4:], style="Heading 2")
            paragraph_keep_with_next(p)
        elif stripped.startswith("## "):
            number_list_id = None
            p = doc.add_paragraph(stripped[3:], style="Heading 1")
            paragraph_keep_with_next(p)
            if stripped == "## 4. Initial findings on the compression–fidelity balance" and not inserted_chart:
                pic = doc.add_picture(str(CHART), width=Inches(6.2))
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(10)
                set_font(cap.add_run("Figure 2. Weighted token reduction in the matched benchmark cohorts."), size=9, color=MUTED, italic=True)
                inserted_chart = True
        elif stripped.startswith("# "):
            number_list_id = None
            continue
        elif stripped == "[CLIENT_SCALE_FIGURE]":
            number_list_id = None
            if CLIENT_SCALE_IMAGE.exists():
                pic = doc.add_picture(str(CLIENT_SCALE_IMAGE), width=Inches(3.15))
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(10)
                set_font(cap.add_run("Figure 1. One client's input-token volume: 2.3B total across 85 active days; 1.4B on the peak day."), size=9, color=MUTED, italic=True)
        elif stripped == "[PAGE_BREAK]":
            number_list_id = None
            doc.add_page_break()
        elif re.match(r"^\d+\. \*\*", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            if number_list_id is None:
                number_list_id = new_numbering_instance(doc, start=1)
            p = doc.add_paragraph(style="List Number")
            apply_numbering(p, number_list_id)
            add_rich_text(p, text)
        elif stripped.startswith("- "):
            number_list_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:])
        elif stripped.startswith("> "):
            number_list_id = None
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.25
            set_font(p.add_run(stripped[2:]), size=12, color=DARK_BLUE, bold=True, italic=True)
        else:
            number_list_id = None
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_rich_text(p, stripped)
    if table_lines:
        add_table(doc, table_lines)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
