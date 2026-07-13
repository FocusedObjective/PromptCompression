import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "json_integrity_before_after.md"
OUTPUT = ROOT / "JSON_Protection_Before_After_Integrity_Study.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203748"
MUTED = "666666"
FILL = "F2F4F7"


def font(run, size=11, color="000000", bold=False, italic=False, name="Calibri"):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def rich(paragraph, text, size=11):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            font(paragraph.add_run(text[cursor:match.start()]), size=size)
        token = match.group(0)
        if token.startswith("**"):
            font(paragraph.add_run(token[2:-2]), size=size, bold=True)
        elif token.startswith("`"):
            font(paragraph.add_run(token[1:-1]), size=size-0.8, name="Consolas", color=DARK_BLUE)
        else:
            font(paragraph.add_run(token[1:-1]), size=size, italic=True)
        cursor = match.end()
    if cursor < len(text):
        font(paragraph.add_run(text[cursor:]), size=size)


def shade(cell, fill=FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    total = sum(widths)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)


def add_table(doc, source_lines):
    data = [[part.strip() for part in line.strip().strip("|").split("|")] for line in source_lines]
    data = [data[0]] + data[2:]
    cols = len(data[0])
    table = doc.add_table(rows=len(data), cols=cols)
    table.style = "Table Grid"
    table.autofit = False
    widths = [3500, 1950, 1950, 1960] if cols == 4 else [4000, 2680, 2680]
    table_geometry(table, widths)
    header_prop = OxmlElement("w:tblHeader")
    header_prop.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header_prop)
    for row_index, values in enumerate(data):
        for col_index, value in enumerate(values):
            cell = table.rows[row_index].cells[col_index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            rich(p, value, size=9.2)
            if row_index == 0:
                shade(cell)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    for kind, value in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
        if kind:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        else:
            node = OxmlElement("w:instrText" if value == " PAGE " else "w:t")
            node.text = value
        run._r.append(node)
    font(run, size=9, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    font(header.add_run("FOCUSEDOBJECTIVE  /  RESEARCH NOTE"), size=8.5, color=MUTED, bold=True)
    page_field(section.footer.paragraphs[0])


def title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run("RESEARCH NOTE"), size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run("Protecting Structured Context"), size=23, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    font(p.add_run("An Initial Before-and-After Study of JSON Integrity in Prompt Compression"), size=14, color=DARK_BLUE)
    for label, value in (
        ("Project", "FocusedObjective / PromptCompression"),
        ("Date", "July 12, 2026"),
        ("Evidence", "39 matched FocusFit prompt pairs"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        font(p.add_run(label + ": "), bold=True)
        font(p.add_run(value))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)


def build():
    doc = Document()
    configure(doc)
    title_block(doc)
    doc.core_properties.title = "Protecting Structured Context"
    doc.core_properties.subject = "Before-and-after JSON integrity study"
    doc.core_properties.author = "FocusedObjective"

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = False
    table_lines = []
    for line in lines:
        if not start:
            if line.strip() == "## Abstract":
                start = True
            else:
                continue
        if table_lines and not line.startswith("|"):
            add_table(doc, table_lines)
            table_lines = []
        if line.startswith("|"):
            table_lines.append(line)
            continue
        text = line.strip()
        if not text:
            continue
        if text.startswith("### "):
            doc.add_paragraph(text[4:], style="Heading 2")
        elif text.startswith("## "):
            doc.add_paragraph(text[3:], style="Heading 1")
        elif re.match(r"^\d+\. \*\*", text):
            p = doc.add_paragraph(style="List Number")
            rich(p, re.sub(r"^\d+\. ", "", text))
        elif text.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            rich(p, text[2:])
        else:
            p = doc.add_paragraph()
            rich(p, text)
    if table_lines:
        add_table(doc, table_lines)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
